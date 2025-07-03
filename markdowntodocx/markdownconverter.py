import re
import io
import requests
import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.constants import CONTENT_TYPE as CT
import random
from docx.shared import Cm
from docx.enum.table import WD_ALIGN_VERTICAL # pylint: disable=no-name-in-module
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.parser import OxmlElement
from docx.exceptions import InvalidXmlError
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml.text.run import CT_R
from docx.oxml.text.hyperlink import CT_Hyperlink
#### Code rewritten and adapted to handle footnotes from baloo-docx ####
from docx.opc.part import PartFactory
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.shared import RGBColor
from docx.oxml.simpletypes import ST_DecimalNumber, ST_String
from docx.opc.constants import NAMESPACE
from docx.oxml.xmlchemy import (
    BaseOxmlElement, RequiredAttribute, ZeroOrMore, ZeroOrOne
)
from docx.shared import Parented

LIMITE_ITERATIONS=10

class Footnote(Parented):
    def __init__(self, f, parent):
        super().__init__(parent)
        self._fn = self._element = self.element = f

class CT_Footnotes(BaseOxmlElement):
    """
    A ``<w:footnotes>`` element, a container for Footnotes properties 
    """

    footnote = ZeroOrMore ('w:footnote', successors=('w:footnotes',))

    @property
    def _next_id(self):
        ids = self.xpath('./w:footnote/@w:id')

        return int(ids[-1]) + 1
    
    def add_footnote(self):
        _next_id = self._next_id
        footnote = CT_Footnote.new(_next_id)
        footnote = self._insert_footnote(footnote)
        return footnote

    def get_footnote_by_id(self, _id):
        namesapce = NAMESPACE().WML_MAIN
        for fn in self.findall('.//w:footnote', {'w':namesapce}):
            if fn._id == _id:
                return fn
        return None
        
class CT_Footnote(BaseOxmlElement):
    """
    A ``<w:footnote>`` element, a container for Footnote properties 
    """
    _id = RequiredAttribute('w:id', ST_DecimalNumber)
    p = ZeroOrOne('w:p', successors=('w:footnote',))

    @classmethod
    def new(cls, _id):
        footnote = OxmlElement('w:footnote')
        footnote._id = _id
        return footnote
    
    def _add_p(self, text):
        _p = OxmlElement('w:p')
        pPr = _p.get_or_add_pPr()
        rstyle = pPr.get_or_add_pStyle()
        rstyle.val = 'FootnoteText'
        
        _r = _p.add_r()
        rPr = _r.get_or_add_rPr()
        rstyle = rPr.get_or_add_rStyle()
        rstyle.val = 'FootnoteReference'
        ref = OxmlElement('w:footnoteRef')
        _r.append(ref)
        _r = _p.add_r()
        ref = OxmlElement('w:footnoteRef')
        _r.append(ref)
        
        run = Run(_r, self)
        run.text = text
        
        self._insert_p(_p)
        return _p


    def _add_p_with_paragraph(self, para):
        _p = para._p
        # paragraph footnote style
        pPr = _p.get_or_add_pPr()
        rstyle = pPr.get_or_add_pStyle()
        rstyle.val = 'FootnoteText'
        para.style = styles["footnote text"]
        # run style (with id of run)
        new_run_element = _p._new_r()
        para.all_runs[0]._element.addprevious(new_run_element)
        rPr = new_run_element.get_or_add_rPr()
        rstyle = rPr.get_or_add_rStyle()
        rstyle.val = 'FootnoteReference'
        r = Run(new_run_element, para)
        r.style = styles["footnote reference"]
        ref = OxmlElement('w:footnoteRef')
        new_run_element.append(ref)
        self._insert_p(_p)
        return _p
    
    @property
    def paragraph(self):
        return DocxParagraph(self.p, self)
    
class CT_FNR(BaseOxmlElement):
    _id = RequiredAttribute('w:id', ST_DecimalNumber)

    @classmethod
    def new (cls, _id):
        footnoteReference = OxmlElement('w:footnoteReference')
        footnoteReference._id = _id
        return footnoteReference

# class CT_Hyperlink(BaseOxmlElement):
#     @classmethod
#     def new (cls):
#         ref = OxmlElement('w:hyperlink')
#         return ref      

class CT_FootnoteRef (BaseOxmlElement):
    @classmethod
    def new (cls):
        ref = OxmlElement('w:footnoteRef')
        return ref       
class FootnotesPart(XmlPart):
    """
    Definition of Footnotes Part
    """
    @classmethod
    def default(cls, package):
        partname = PackURI("/word/footnotes.xml")
        content_type = CT.WML_FOOTNOTES
        element = parse_xml(cls._default_footnotes_xml())
        return cls(partname, content_type, element, package)


docx.oxml.register_element_cls('w:footnotes', CT_Footnotes)
docx.oxml.register_element_cls('w:footnote', CT_Footnote)
docx.oxml.register_element_cls('w:footnoteReference', CT_FNR)
docx.oxml.register_element_cls('w:footnoteRef', CT_FootnoteRef)
#docx.oxml.register_element_cls('w:hyperlink', CT_Hyperlink)
PartFactory.part_type_for[CT.WML_FOOTNOTES] = FootnotesPart
##### END OF FOOTNOTES CODE ####
#### EXTEND PARAGRAPH TO BE ABLE TO READ HYPERLINKS ####
class DocxParagraph(docx.text.paragraph.Paragraph):
    def __init__(self, *args, **kwargs):
        super().__init__( *args, **kwargs)
    
    def remove(self, element):
        try:
            self._p.remove(element)
        except ValueError:
            for elem in self._p:
                if isinstance(elem, CT_Hyperlink):
                    try:
                        elem.remove(element)
                    except ValueError:
                        pass

    def index(self, element):
        try:
            return self._p.index(element)
        except ValueError:
            for i, elem in enumerate(self._p):
                if isinstance(elem, CT_Hyperlink):
                    if element in elem:
                        return i

    def get_all_text(self):
        """
        String formed by concatenating the text of each run in the paragraph.
        Tabs and line breaks in the XML are mapped to ``\\t`` and ``\\n``
        characters respectively.

        Assigning text to this property causes all existing paragraph content
        to be replaced with a single run containing the assigned text.
        A ``\\t`` character in the text is mapped to a ``<w:tab/>`` element
        and each ``\\n`` or ``\\r`` character is mapped to a line break.
        Paragraph-level formatting, such as style, is preserved. All
        run-level formatting, such as bold or italic, is removed.
        """
        text = ''
        for run in self.all_runs:
            text += run.text
        return text

    @property
    def all_runs(self):
        runs = []
        for _elem in self._p:
            if isinstance(_elem, CT_R):
                runs.append(Run(_elem, self))
            if isinstance(_elem, CT_Hyperlink):
                for _r in _elem:
                    runs.append(Run(_r, _elem))
        return runs
####
footnotes = {}
default_styles_names = {
        "Hyperlink": ("Hyperlink",),
        "Code": ("Code", "macro"),
        "Code Car": ("CodeStyle", "Code Car", "Macro Text Char"),
        "BulletList": ("BulletList", "List Paragraph"),
        "Cell": ("Cell", "No Spacing"),
        "Header1": ("Heading 1", "Header"),
        "Header2": ("Heading 2", "Header"),
        "Header3": ("Heading 3", "Header"),
        "Header4": ("Heading 4", "Header"),
        "Header5": ("Heading 5", "Header"),
        "Header6": ("Heading 6", "Header"),
        "Table": ("Table", "Table Grid"),
    }
styles = {}
code_style = None
hyperlink_style = None

def convertMarkdownInFile(infile, outfile, styles_names=None):
    global default_styles_names
    global styles
    if styles_names:
        for key, val in styles_names.items():
            default_styles_names[key] = val
    document = docx.Document(infile)
    for style in document.styles:
        styles[style.name] = style
    for key, style_name in default_styles_names.items():
        if isinstance(style_name, tuple):
            found = False
            for name in style_name:
                if name in styles:
                    default_styles_names[key] = name
                    found = True
                    break
            if not found and name != "Hyperlink":
                return False, "Error in template. There is no "+str(name)+ " style in the doc. Searched for : "+str(key)
        else:
            if style_name not in styles and name != "Hyperlink":
                return False, "Error in template. There is a style missing : "+str(style_name)
    global code_style
    global hyperlink_style
   
    code_style = styles.get(default_styles_names.get("Code Car", "Code Car"), "macro")
    hyperlink_style = styles.get(default_styles_names.get("Hyperlink", "Hyperlink"), None)
    markdownToWordInDocument(document)
    document.save(outfile)
    return True, outfile
    
def markdownToWordInDocument(document):
    ps = [ps for ps in getParagraphs(document)]
    state = "normal"
    paragraph_i = 0
    while paragraph_i < len(ps):
        paragraph = ps[paragraph_i]
        state, deleted_count = markdownToWordInParagraph(document, paragraph, state)
        if deleted_count > 0:
            paragraph_i += deleted_count
        paragraph_i += 1
    ps = getParagraphs(document)
    for paragraph in ps:
        state = markdownToWordInParagraphCar(document, paragraph, state)
    

def markdownToWordFromString(string, outfile):
    document = docx.Document()
    paragraphs = string.replace("\r","").split("\n")
    for para in paragraphs:
        document.add_paragraph(para)
    document.save(outfile)
    return convertMarkdownInFile(outfile, outfile)
    
def getParagraphs(document):
    """ Retourne un generateur pour tous les paragraphes du document.
        La page d'entête n'étant pas incluse dans documents.paragraphs."""
    body = document._body._body # pylint: disable=protected-access
    ps = body.xpath('//w:p')
    for p in ps:
        yield DocxParagraph(p, document._body) # pylint: disable=protected-access

def split_run_in_two(paragraph, run, split_index):
    index_in_paragraph = paragraph.index(run.element) # pylint: disable=protected-access
    text_before_split = run.text[0:split_index]
    text_after_split = run.text[split_index:]
    run.text = text_before_split
    new_run = paragraph.add_run(text_after_split)
    copy_format_manual(run, new_run)
    paragraph._p[index_in_paragraph+1:index_in_paragraph+1] = [new_run.element] # pylint: disable=protected-access
    return [run, new_run]

def split_run_in_three(paragraph, run, split_start, split_end):
    first_split = split_run_in_two(paragraph, run, split_end)
    second_split = split_run_in_two(paragraph, run, split_start)
    return second_split + [first_split[-1]]

def copy_format_manual(runA, runB):
    fontB = runB.font
    fontA = runA.font
    fontB.bold = fontA.bold
    fontB.italic = fontA.italic
    fontB.underline = fontA.underline
    fontB.strike = fontA.strike
    fontB.subscript = fontA.subscript
    fontB.superscript = fontA.superscript
    fontB.size = fontA.size
    try:
        fontB.highlight_color = fontA.highlight_color
    except InvalidXmlError as e:
        pass
    except ValueError as e:
        pass
    fontB.color.rgb = fontA.color.rgb

def get_next_paragraph(paragraph):
    try:
        return DocxParagraph(paragraph._p.getnext(), paragraph._parent) # pylint: disable=protected-access
    except AttributeError:
        return None


def markdownArrayToWordList(document, paragraph, state):
    if paragraph.text.strip() == "":
        return state, 0
    table_line_regex = re.compile(r"^\s*(?:\|[^|\n]*)+$", re.MULTILINE)
    next_para = get_next_paragraph(paragraph)
    matched = []
    matched_para = []
    matched_line = re.findall(table_line_regex, paragraph.text)
    if matched_line:
        matched_para.append(paragraph)
    while matched_line:
        matched.extend(matched_line)
        if next_para is None:
            break
        matched_line = re.findall(table_line_regex, next_para.text)
        if matched_line:
            matched_para.append(next_para)
        next_para = get_next_paragraph(next_para)
        
        if next_para is None or next_para.text is None or next_para.text.strip() == "":
            if matched_line:
                matched.extend(matched_line)
            break
    if len(matched) == 0:
        return state, 0
    
    nb_columns = len([x.strip() for x in matched[0].strip().split("|") if x.strip() != ""])
    if nb_columns < 1 or len(matched) < 3:
        return state, 0
    
    array = document.add_table(rows=len(matched) - 1, cols=nb_columns) # remove header/body line separator
    try:
        array.style = styles[default_styles_names.get("Table", None)]
    except KeyError:
        pass
    horizontal_alignment = [None] * nb_columns
    for i_row, match in enumerate(matched):
        line = match.strip()
        columns = [x.strip() for x in line.split("|") if x.strip() != ""]
        # if len(columns) != nb_columns:
        #     raise ValueError("The array with following headers : "+str(matched[0])+" is supposed to have "+str(nb_columns)+ \
        #                         " columns but the line "+str(line)+" has "+str(len(columns))+" columns")
        for i_column, column in enumerate(columns):
            if i_column >= nb_columns:
                break
            if i_row == 1:
                if column.strip().startswith(":") and not column.strip().endswith(":"):
                    horizontal_alignment[i_column] = WD_ALIGN_PARAGRAPH.LEFT
                elif not column.strip().startswith(":") and column.strip().endswith(":"):
                    horizontal_alignment[i_column] = WD_ALIGN_PARAGRAPH.RIGHT
                elif column.strip().startswith(":") and column.strip().endswith(":"):
                    horizontal_alignment[i_column] = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cell = array.cell(0 if i_row == 0 else i_row-1, i_column)
                fill_cell(document, cell, column, horizontal_align=horizontal_alignment[i_column])
    move_table_after(array, paragraph)
    for para in matched_para:
        delete_paragraph(para)
    #delete_paragraph(paragraph)
    return state, len(matched_para)

def markdownUnorderedListToWordList(paragraph, style, state):
    regex = re.compile(r"^\s*[\*|\-|\+]\s([^\n]+)", re.MULTILINE)
    matched = re.findall(regex, paragraph.text)
    if len(matched) > 0:
        start = paragraph.text.index(matched[0])
        try:
            end = paragraph.text.index(matched[-1], start+len(matched[0]))+len(matched[-1])
            text_end = paragraph.text[end:]
        except ValueError:
            text_end = ""
        paragraph.text = paragraph.text[:start-2].strip() # -2 for list marker + space
        new_p = paragraph # init for loop
        for match in matched:
            new_p = insert_paragraph_after(new_p)
            new_p.style = styles[default_styles_names.get("BulletList","BulletList")]
            r = new_p.add_run()
            r.add_text(match)
        if text_end.strip() != "":
            insert_paragraph_after(new_p, text_end)
        if paragraph.text.strip() == "":
            delete_paragraph(paragraph)
    return state

def mardownCodeBlockToWordStyle(paragraph, code_style, state):
    if state == "code_block":
        paragraph.style = code_style
    if "```" in paragraph.text and state != "code_block":
        state = "code_block"
        text_bits = paragraph.text.split("```")
        paragraph.text = text_bits[0].strip()
        paragraph = insert_paragraph_after(paragraph, "```".join(text_bits[1:]), code_style)

    
    if "```" in paragraph.text and state == "code_block":
        state = "normal"
        text_bits = paragraph.text.split("```")
        paragraph.text = text_bits[0].strip()
        insert_paragraph_after(paragraph, "```".join(text_bits[1:]))
        
    return state

def markdownToWordInParagraph(document, paragraph, state):
    deleted_count = 0
    state, deleted_count = markdownArrayToWordList(document, paragraph, state)
    state = markdownUnorderedListToWordList(paragraph, styles[default_styles_names.get("BulletList","BulletList")], state)
    state = mardownCodeBlockToWordStyle(paragraph, styles.get(default_styles_names.get("Code","Code"), "Macro Text"), state)
    return state, deleted_count

def setColor(paragraph, run, match):
    # match.group(0) is <color:RGB>text</color> 
    # match.group(1) is <color
    # match.group(2) is RGB>text
    # match.group(3) is </color>
    #or
    # match.group(0) is <span style="color: rgb(0,0,0);">text</span>
    # match.group(1) is <span style="color: 
    # match.group(2) is rgb(0,0,0);">text
    # match.group(3) is </span>
    

    splitted = match.group(2).split(">")
    if splitted[0].startswith("rgb("):
        r,g,b = splitted[0][4:-1].split(",")
        b = b.split(")")[0]
        try:
            run.font.color.rgb = RGBColor(int(r.strip()), int(g.strip()), int(b.strip()))
        except ValueError:
            pass
    else:
        color_string = splitted[0]
        color_string = color_string.strip().replace("#", "")
        if len(color_string) != 6:
            raise ValueError("RGB hex string must be exactly 6 characters long")
        run.font.color.rgb = RGBColor.from_string(color_string)
    initial_len = len(run.text)
    run.text = ">".join(splitted[1:])
    return initial_len-len(run.text), False, "normal"

def markdownToWordInParagraphCar(document, paragraph, state):
    if paragraph.style.name == code_style.name or paragraph.style.name == "Code":
        return state
    for i in range(LIMITE_ITERATIONS):
        original_text = paragraph.text
        markdownHeaderToWordStyle(paragraph)
        transform_marker(paragraph, "==", setHighlight)
        transform_marker(paragraph, "**", setBold)
        transform_marker(paragraph, "__", setBold)
        transform_marker(paragraph, "*", setItalic)
        transform_marker(paragraph, "_", setItalic)
        transform_marker(paragraph, "~~", setStrike)
        transform_marker(paragraph, "`", setCode)
        
        transform_regex(paragraph, r"(<color:)(.*?>.*?)(</color>)", (delCar, setColor, delCar))
        transform_regex(paragraph, r"(<span\s+style=\"color: )(.*?>.*?)(</span>)", (delCar, setColor, delCar))

        #bookmarks [#bookmark]
        lambda_book = lambda para, run, match: setBookmark(document, para, run, match)
        transform_regex(paragraph, r"(\[#)([^\]\n]*)(\])(?!\w)", (delCar, lambda_book, delCar))
        # markdown hyper link in the format [text to display](link)
        transform_regex(paragraph, r"(?<!\!)(\[)([^\]|^\n]+)(\]\()([^\)|^\n]+)(\))", (delCar, setHyperlink, delCar, delCar, delCar))
        # markdown image hyper link to incorporate in the format ![alt text to display](link)
        transform_regex(paragraph, r"(\!\[)([^\]|^\n]+)(\]\()([^\)|^\n]+\.(?:png|jpg|jpeg|gif))(\))", (delCar, linkImageToImage, delCar, delCar, delCar))
        # just make left hyperlinks clickable
        #LAST IS footnotes BECAUSE THE PARAGRAPH IS MOVED
        #inline footnotes ^[footnote text]
        
        lambda_foot = lambda para, run, match: setInlineFootnote(document, para, run, match)
        transform_regex(paragraph, r"(\^\[)([^\]\n]*)(\])(?!\w)", (delCar, lambda_foot, delCar))
        # footnotes insertion [^footnote id name]
        lambda_declare_foot = lambda para, run, match: declareFootnote(document, para, run, match)
        transform_regex(paragraph, r"(\[\^)([^\]\n]*)(\])(?!\w|:)", (delCar, lambda_declare_foot, delCar))
        # footnotes text description [^footnote id name]: indented text with possibly many paragraphs
        if state.startswith("inFootnoteDefinition:"):
            if paragraph.text.strip() == "":
                state = "normal"
            else:
                footnote_id = state.split(":")[1]
                footnote = footnotes[footnote_id]
                pPr = paragraph._p.get_or_add_pPr()
                rstyle = pPr.get_or_add_pStyle()
                rstyle.val = 'FootnoteText'
                paragraph.style = styles["footnote text"]
                footnote._fn._insert_p(paragraph._p)
        else:
            state = transform_regex(paragraph, r"^(\[\^)([^\]\n]*)(\]:)(?!\w)(.+(?:\n[ \t]+.+)*)", (delCar, delCar, delCar, defineFootnote))
        # hyperlinks
        if state.startswith("inFootnoteDefinition"):
            footnote_id = state.split(":")[1]
            footnote = footnotes[footnote_id]
            paragraph = DocxParagraph(paragraph._p, footnote)
            lambda_sethyperlink_footnote = lambda para, run, match: setHyperlink(para, run, match, document=document, is_footnote=True)
            transform_regex(paragraph, r"(https?:\/\/(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&//=]*[-a-zA-Z0-9@:%_\+.~#?&//=]))", (lambda_sethyperlink_footnote,))
        else:
            transform_regex(paragraph, r"(https?:\/\/(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&//=]*[-a-zA-Z0-9@:%_\+.~#?&//=]))", (setHyperlink,))
        if original_text == paragraph.text:
            break
        else:
            continue
    return state

def setHyperlink(paragraph, run, match, **kwargs):
    # Handle if already a link
    if hasattr(run._parent,"tag") and run._parent.tag.endswith("hyperlink"):
        return 0, False, "normal"
    if run.style.style_id == code_style.style_id:
        return 0, False, "normal"
    run.font.underline = True
    if hyperlink_style is not None:
        run.style = hyperlink_style
    else:
        run.font.color.rgb = RGBColor.from_string("0000FF")
    # link are not counted as deleted anymore // deletedCars = len(run.text)
    try:
        link_text = match.group(2)
        link_url = match.group(4)
    except:
        link_text = kwargs.get("text", match.group(0))
        link_url = kwargs.get("url", match.group(0))

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    external =  link_url.startswith("http")
    if external:
        part = paragraph.part
        if kwargs.get("is_footnote", False):
            part = _footnotes_part = kwargs["document"]._part.part_related_by(RT.FOOTNOTES)
        r_id = part.relate_to(link_url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    else:
        hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_url)
    index_in_paragraph = paragraph._p.index(run.element)
    run.element.text = link_text
    hyperlink.append(run.element)
    paragraph._p[index_in_paragraph:index_in_paragraph] = [hyperlink]
    # Delete this if using a template that has the hyperlink style in it
     # link are not counted as deleted anymore //return 0, True, "normal"
    return 0, False, "normal"

def add_footnote(document):
    _footnotes_part = document._part.part_related_by(RT.FOOTNOTES)
    footnotes_part = _footnotes_part.element
    footnote = footnotes_part.add_footnote()

    footnote = Footnote(footnote, document._part)
    return footnote

def add_footnote_reference(run, footnote_element):
    rPr = run._r.get_or_add_rPr()
    rstyle = rPr.get_or_add_rStyle()
    rstyle.val = 'FootnoteReference'
    run.style = styles["footnote reference"]
    reference = OxmlElement('w:footnoteReference')
    reference._id = footnote_element._id
    run._r.append(reference)

def setInlineFootnote(document, paragraph, run, match):
    """Set the inline footnote.
    run: document run that will be transformed to a footnote."""
    deletedCars = len(run.text)
    run.text= "" # Run text is removed to be placed inside the footnote paragraph
    #footnotes
    footnote = add_footnote(document) # create the footnote section for the document (empty)
    gr = re.search(r"(https?:\/\/(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&//=]*[-a-zA-Z0-9@:%_\+.~#?&//=]))", match.group(2))
    _p = footnote._fn._add_p(" ") # create a paragraph with a run containing a space to separate the footnote id from the text
    para = DocxParagraph(_p, footnote)
    para.style = styles["footnote text"] # set footnote style
    para.add_run(str(match.group(2)))   # put match text inside the footnotes
    if gr is not None: # is link
        h_deletedCars, h_deletedRun, h_state = setHyperlink(para, para.all_runs[-1], gr, text=str(match.group(2)), is_footnote=True, document=document)

    para.all_runs[0].style = styles["footnote reference"] # set footnote id style
    # footnotes reference
    add_footnote_reference(run, footnote._fn)
    
    return deletedCars, False, "normal"

def declareFootnote(document, paragraph, run, match):
    deletedCars = len(run.text)
    run.text = ""
    #footnotes
    footnote = add_footnote(document)
    
    global footnotes
    footnotes[match.group(2)] = footnote
    add_footnote_reference(run, footnote._fn)

    return deletedCars, False, "normal"

def defineFootnote(paragraph, run, match):
    #footnotes
    footnote_id = match.group(2)
    footnote = footnotes[footnote_id]
    
    _p = footnote._fn._add_p_with_paragraph(paragraph)
    return 0, False, "inFootnoteDefinition:"+str(footnote_id)


def setBookmark(document, paragraph, run, match):
    """Set the bookmark
    function adapted from https://stackoverflow.com/questions/57586400/how-to-create-bookmarks-in-a-word-document-then-create-internal-hyperlinks-to-t
    """
    tag = run._r
    deleted = len(run.text)
    run.text = ""
    body = document._body._body
    ids = body.xpath('//w:bookmarkStart/@w:id')
    if ids:
        id = int(ids[-1]) + 1
    else:
        id = 1
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), str(id))
    start.set(docx.oxml.ns.qn('w:name'), match.group(2))
    tag.append(start)
    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), str(id))
    end.set(docx.oxml.ns.qn('w:name'), match.group(2))
    tag.append(end)
    return deleted, False, "normal"

def linkImageToImage(para, run, match):
    link_url = match.group(4)
    if link_url.strip().startswith("file://"):
        path = link_url.strip()[7:]
        try:
            with open(path, "rb") as f:
                data = io.BytesIO(f.read())
        except Exception as e:
            print("Error opening file: ", e)
            data = None
    else:
        data = downloadImgData(link_url)
    if data is not None:
        text_len = len(run.text)
        run.text = ""
        run.add_picture(data, width=Cm(17.19))
        return text_len, False, "normal"
    return 0, False, "normal"

def setHighlight(para, run, match):
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return 0, False, "normal"

def setBold(para, run, match):
    run.bold = True
    return 0, False, "normal"

def setItalic(para, run, match):
    run.italic = True
    return 0, False, "normal"

def setStrike(para, run, match):
    run.font.strike = True
    return 0, False, "normal"

def setCode(para, run, match):
    try:
        run.style = code_style
    except ValueError as e:
        raise ValueError(f"Style for code is {code_style.name} and is of type PARAGRAPH WHERE A CHARACTER STYLE IS NEEDED") from e
    
    return 0, False, "normal"

def delCar(para, run, match):
    ret = len(run.text)
    run.text = ""
    para.remove(run.element)
    return ret, True, "normal"

def transform_marker(paragraph, marker, func,content_regex=None):
    if content_regex is None:
        content_regex = r"[^"+re.escape(marker[0])+r"\n]*"
    marker = re.escape(marker)
    
    regex = r"(?<!\w)"+"("+marker+")("+content_regex+")("+marker+")"+r"(?!\w)"
    return transform_regex(paragraph, regex, (delCar, func, delCar))


def transform_regex(paragraph, regex, funcs):
    deletedCars = 0
    state = "normal"
    regex = re.compile(regex, re.MULTILINE)
    # find every iteration of marker+content+marker in paragraph
    for match in regex.finditer(paragraph.text):
        # get starting marker run index and ending marker run index
        positions = []
        core_pos = [x[0]-deletedCars for x in match.regs[1:]]
        positions += core_pos # Get starting pos of match of each group
        positions.append(match.regs[0][1]-deletedCars)
        runs = getRunsIndexFromPositions(paragraph, positions)
        
        # merge non-contiuous run matched
        # Check if last run has at least one character to be merged. (the last pos is the end of the match
        pos = ([x[0] for x in runs if x is not None])
        start = min(pos)
        end = max(pos)
        while start < end:
            paragraph.all_runs[end-1].text += paragraph.all_runs[end].text
            paragraph.all_runs[end].text = ""
            paragraph.remove(paragraph.all_runs[end]._element)
            end -= 1
        # find marker position in run and split
        runs = getRunsIndexFromPositions(paragraph, positions)
        prev = None
        for run_pos in runs[::-1]:
            if run_pos is None:
                prev = [None, 0] # force split
                continue
            # Split runs if needed
            run = paragraph.all_runs[run_pos[0]]
            if run_pos[1] != 0: # if not at the beginning of the run
                split_run_in_two(paragraph, run, run_pos[1])
            prev = run_pos
        runs = getRunsIndexFromPositions(paragraph, core_pos)
        # apply transformation func on all runs found
        deleted_runs = 0
        for i, func in enumerate(funcs):
            run = paragraph.all_runs[runs[i][0] - deleted_runs]
            deleted_count, deleted_run, state = func(paragraph, run, match)
            deleted_runs += 1 if deleted_run else 0
            deletedCars += deleted_count
    return state

def markdownHeaderToWordStyle(paragraph):
    for match in re.finditer(r"^[ \t]*(#{1,6}) (.+)\s*", paragraph.text, re.MULTILINE):
        rest = paragraph.text.replace(match.group(0), "").strip()
        if rest != "":
            insert_paragraph_after(paragraph, rest)
        paragraph.text = re.sub(r"^[ \t]*#{1,6} (.+)\s*", r"\1", match.group(0)).strip()
        headersize = len(match.group(1))
        header_style = styles.get(default_styles_names.get("Header"+str(headersize), "Heading "+str(headersize)), None)
        if header_style is not None:
            paragraph.style = header_style
        

def getRunsIndexFromPositions(paragraph, positions):
    r"""
    Returns a list of tuples (runIndex, positionInRun) for each caracter position in the list positions
    example:
        paragraph.text is "This is an *example* of a paragraph"
        runs look like this:
            [
                "This is an *e",
                "xample",
                "* of a pa",
                "ragraph"
            ]
        regex match on paragraph for \*example\* is (11, 20) 
        we want to know which run contains the 11th caracter  and the 20th caracter
        the 11th is in the run 0 at position 11 and the 20th is in the run 2 at position 1 (it's the space)
        so we return [(0,11), (2, 1)],
        2nd example:
        paragraph.text is "**This is bold text**"
        we want to match **.+**
        runs look like this:
        [
            "**This is bold text**"
        ]
        returns [(0, 0), (0, 20)]
        """
    countedLetters = 0
    prevCountedLetters = 0
    ret = [None] * len(positions)
    for i, run in enumerate(paragraph.all_runs):
        prevCountedLetters = countedLetters
        countedLetters += len(run.text)
        for j, pos in enumerate(positions):
            if prevCountedLetters <= pos and pos < countedLetters:
                ret[j] = (i,pos-prevCountedLetters)
    return ret

def fill_cell(document, cell, text, font_color=None, bg_color=None, bold=False, horizontal_align=None):
    """
    Fill a table's cell's background with a background color, a text and a font_color for this text
    Also sets the vertical alignement of every cell as centered.
        Args:
            cell: the cell we want to fill
            text: the text to be written inside the cell
        Optional Args:
            font_color: a new font color to use for this text at RGB format (docx rgb). Default is None
            bg_color: A backgroud color to use for the cell at hexa rgb format (FFFFFF is white) default is None.
                        The color is written in xml directly as python-docx does not give a function to do that.
    """
    while len(cell.paragraphs) > 0:
        delete_paragraph(cell.paragraphs[0])
    p = cell.add_paragraph(text)
    p = DocxParagraph(p._element, cell) # pylint: disable=protected-access
    p.style = styles["Cell"]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p.alignment = horizontal_align
    if p.all_runs:
        p.all_runs[0].bold = bold
        if font_color is not None:
            p.all_runs[0].font.color.rgb = font_color
    if bg_color is not None:
        shading_elm_1 = parse_xml((r'<w:shd {} w:fill="'+bg_color+r'"/>').format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)  # pylint: disable=protected-access


def insertPageBreak(paragraph):
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

def move_table_after(table, paragraph):
    """
    Move a given table after a given paragraph
        Args:
            table: the table to move
            paragraph: the paragraph to put the table after.
    """
    tbl, p = table._tbl, paragraph._p  # pylint: disable=protected-access
    p.addnext(tbl)

def delete_paragraph(paragraph):
    """
    Delete a paragraph.
        Args:
            paragraph: the paragraph object to delete.
    """
    p = paragraph._element  # pylint: disable=protected-access
    try:
        p.getparent().remove(p)
    except Exception: # pylint: disable=broad-except
        print("No parent found for element "+str(p))
    p._p = p._element = None  # pylint: disable=protected-access


def downloadImgData(url):
    try:
        data = requests.get(url, timeout=3)
        if data.status_code != 200:
            return None
    except Exception as e:
        return None
    data = data.content
    data = io.BytesIO(data)
    return data

def getParagraphs(document):
    """ Retourne un generateur pour tous les paragraphes du document.
        La page d'entête n'étant pas incluse dans documents.paragraphs."""
    body = document._body._body # pylint: disable=protected-access
    ps = body.xpath('//w:p')
    for p in ps:
        yield DocxParagraph(p, document._body) # pylint: disable=protected-access


def set_hyperlink(paragraph, run, url, text, style):
    # This gets access to the document.xml.rels file and gets a new relation id value
    run.font.underline = True
    if style is not None:
        run.style = style
    else:
        run.font.color.rgb = RGBColor(0, 0, 255)
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    index_in_paragraph = paragraph._p.index(run.element)
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    run.element.text = text
    hyperlink.append(run.element)
    paragraph._p[index_in_paragraph:index_in_paragraph] = [hyperlink]
    # Delete this if using a template that has the hyperlink style in it
    return hyperlink

def insert_paragraph_after(paragraph, text=None, style=None):
    """
    Insert a new paragraph after the given paragraph.
        Args:
            paragraph: the paragraph object after which the new paragraph will be created
        
        Optional Args:
            text: a string of text to write in the new paragraph, Default = None
            style: a style to be applied on the new paragraph, Default = None

        Returns:
            Returns a paragraph object for the new added paragraph
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)  # pylint: disable=protected-access
    new_para = DocxParagraph(new_p, paragraph._parent)  # pylint: disable=protected-access
    if text is not None:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


if __name__ == '__main__':
    res, msg = convertMarkdownInFile("examples/in_document.docx", "examples/out_document.docx" ,{"Header":"Header"})
#     res, msg = markdownToWordFromString("""# H1 Header: Welcome to My Markdown Guide!

# ## H2 Header: Quick Overview
# Markdown is a lightweight markup language that you can use to add formatting elements to plaintext text documents. Created by John Gruber in 2004, Markdown is now one of the world’s most popular markup languages.

# ### H3 Header: What Markdown Can Do

# #### H4 Header: Formatting Text
# You can do numerous things with Markdown format, including:

# - **Bold** text
# - *Italic* text
# - **_Combined emphasis_**

# > **Note:** Markdown is not the same as Markup. They’re different, remember!

# #### H4 Header: Creating Lists

# ##### H5 Header: Unordered Lists
# - Item one
# - Item two
#   - Sub Item one
#   - Sub Item two

# ##### H5 Header: Ordered Lists
# 1. First item
# 2. Second item
#     1. Subitem
#     2. Subitem

# #### H4 Header: Adding Links and Images

# Here is a clickable link to [OpenAI](https://www.openai.com), the organization behind Assistant.

# Here is an image:

# [Alt text for image](https://via.placeholder.com/150)

# #### H4 Header: Inserting Code

# ```python
# # This is some Python code
# def say_hello(name):
#     print("Hello, " + name)
# ```

# ##### H5 Header: Inline Code
# You can also use inline `code` within your text.

# ### H3 Header: Using Blockquotes

# > Markdown uses email-style > characters for blockquoting.
# > It’s very handy for email mimicking."""
# , "examples/out_string.docx")
    if res:
        print("Success : output document path is "+msg)
    else:
        print("Error in document : "+msg)